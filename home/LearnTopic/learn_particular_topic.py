import streamlit as st
import google.generativeai as genai
import streamlit.components.v1 as components
import markdown2
import pdfkit
import tempfile
from docx import Document  # For Word file generation
# python-docx



genai.configure(api_key="AIzaSyB7n_R6Bs5JxGyzzS1V4DftNR9-Wu2j-2o")
# Configure generation settings
generation_config = {
    "temperature": 0.3,  # Controls randomness (0 = deterministic, 1 = creative)
    "max_output_tokens": 1200,  # Increase token limit for detailed responses
    "top_p": 0.9,  # Controls diversity via nucleus sampling
    "top_k": 50,  # Limits sampling to the top-k most likely tokens
    "stop_sequences": [
        "References:",
        "End of response",
    ],  # Prevents unnecessary text
}
# Function to convert Markdown to PDF
def convert_markdown_to_pdf(md_text):
    # Convert Markdown to HTML
    html_text = markdown2.markdown(md_text, extras=["tables", "fenced-code-blocks", "strike", "footnotes"])
    
    # Add basic HTML structure with company name as a big title
    html_content = f"""
    <html>
    <head>
        <meta charset='utf-8'>
        <style>
            body {{ font-family: Arial, sans-serif; padding: 20px; }}
            h1 {{ color: #333; font-size: 32px; text-align: center; margin-bottom: 30px; }}
            h2, h3, h4, h5, h6 {{ color: #333; }}
            pre {{ background-color: #f4f4f4; padding: 10px; border-radius: 5px; }}
            code {{ background-color: #f4f4f4; padding: 2px 4px; border-radius: 3px; }}
            blockquote {{ border-left: 4px solid #ccc; padding-left: 10px; margin-left: 0; color: #666; }}
            ul, ol {{ padding-left: 20px; }}
        </style>
    </head>
    <body>
        <h1>Ai-prepMaster</h1>
        {html_text}
    </body>
    </html>
    """
    
    return html_content

# Function to convert Markdown to Word (.docx) with proper styling
def convert_markdown_to_word(md_text):
    try:
        doc = Document()
        doc.add_heading("Ai-prepMaster", level=1)  # Add company name as title

        # Split the text into lines and process each line
        for line in md_text.split("\n"):
            if line.startswith("# "):  # Heading 1
                doc.add_heading(line[2:], level=1)
            elif line.startswith("## "):  # Heading 2
                doc.add_heading(line[3:], level=2)
            elif line.startswith("### "):  # Heading 3
                doc.add_heading(line[4:], level=3)
            else:
                # Handle bold (**text**) and italic (*text*)
                paragraph = doc.add_paragraph()
                parts = line.split("**")
                for i, part in enumerate(parts):
                    if i % 2 == 1:  # Bold text
                        paragraph.add_run(part).bold = True
                    else:
                        subparts = part.split("*")
                        for j, subpart in enumerate(subparts):
                            if j % 2 == 1:  # Italic text
                                paragraph.add_run(subpart).italic = True
                            else:
                                paragraph.add_run(subpart)
        return doc
    except Exception as e:
        st.error(f"Error generating Word document: {e}")
        return None

def generate_response(topic, context=None):
    if context:
        prompt = f"""
            Task: Act as an expert educator and curriculum designer. Continue the discussion based on the previous context and the new topic: {topic}.

            Previous Context:
            {context}

            New Topic:
            {topic}

            Requirements:
            - Provide a detailed response that builds on the previous context.
            - Include practical examples, code snippets (if applicable), and step-by-step explanations.
            - If the topic is related to programming (e.g., Java, Python), provide executable code examples with explanations.
            - If the new topic is unrelated, treat it as a standalone topic but acknowledge the previous context briefly.
        """
    else:
        st.session_state.first_topic = topic  # Save the first topic for file naming
        prompt = f"""
            Task: Act as an expert educator and curriculum designer. Generate a structured, detailed educational overview for the topic: {topic}.

            Requirements:
            - Start with a concise definition and significance of the topic.
            - Provide key concepts, learning objectives, subtopics, applications, and resources.
            - Include practical examples, code snippets (if applicable), and step-by-step explanations.
            - If the topic is related to programming (e.g., Java, Python), provide executable code examples with explanations.
        """
    try:
        model = genai.GenerativeModel("gemini-pro")
        response = model.generate_content(prompt, generation_config=generation_config)
        return response.text
    except Exception as e:
        st.error(f"Error: {e}")
        return None

def learn_particular_topic_interface():
    st.subheader("Learn Particular Topic")
    
    # Initialize session state variables
    if "generated_details" not in st.session_state:
        st.session_state.generated_details = []
    if "communication_history" not in st.session_state:
        st.session_state.communication_history = []
    if "show_input" not in st.session_state:
        st.session_state.show_input = True
    if "first_topic" not in st.session_state:
        st.session_state.first_topic = ""

    

    # Display previous communication history
    if st.session_state.communication_history:
        for entry in st.session_state.communication_history:
            st.write(entry)

    # Input field and buttons at the bottom
    if st.session_state.show_input:
        topic = st.text_input("Enter a topic you want to Learn:")
        if st.button("Learn"):
            if topic:
                with st.spinner("Generating details..."):
                    # Get the previous context (last response) if available
                    context = None
                    if st.session_state.communication_history:
                        context = st.session_state.communication_history[-1]  # Use the last response as context

                    # Generate response based on the topic and context
                    response = generate_response(topic, context)
                    if response:
                        # Add user input and response to communication history
                        st.session_state.communication_history.append(f"**Your input:** {topic}\n\n**Response:**\n{response}\n\n---\n")
                        st.session_state.show_input = False  # Hide input after generating response
                        st.rerun()  # Rerun the app to update the UI
            else:
                st.warning("Please enter a topic to generate a response.")

    # Show next input field and download button after generating details
    if not st.session_state.show_input:
        if st.button("Next Input"):
            st.session_state.show_input = True
            st.rerun()  # Rerun the app to show the input field again

        if st.button("Generate PDF"):
            with st.status("Generating PDF...",expanded=False) as status:
                # Create a PDF from the communication history
                md_text = "\n".join(st.session_state.communication_history) if isinstance(st.session_state.communication_history, list) else str(st.session_state.communication_history)
                html_content = convert_markdown_to_pdf(md_text)

                with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp_pdf:
                    pdfkit.from_string(html_content, tmp_pdf.name)
                    
                    # Provide the download link
                    st.success("PDF generated successfully!")
                    st.download_button(label="Download PDF", data=open(tmp_pdf.name, "rb").read(), file_name=f"{st.session_state.first_topic}.pdf", mime="application/pdf")

                # Save as PDF button
                show_print_button = """
                <script>
                    function print_page(obj) {
                        parent.window.print();
                    }
                </script>
                <style>
                    #save_pdf_button {
                        background-color: rgba(19,23,32,1);
                        color: white;
                        padding: 10px 20px;
                        margin: 0px;
                        font-size: 16px;
                        border: 1px solid rgba(65,68,76,1);
                        border-radius: 5px;
                        cursor: pointer;
                        transition: 0.3s;
                        position: absolute;
                        left : 0;
                        top: 0;
                    }
                    #save_pdf_button:hover {
                        background-color:rgba(19,23,32,1);
                        border: 1px solid red;
                        color: red;
                    }
                </style>

                <button id="save_pdf_button" onclick="print_page(this)">
                    Save current window as pdf
                </button>
                """
                components.html(show_print_button)
            status.update(label="Processing complete!", state="complete", expanded=True)  # Finish status

        if st.button("Generate Word Document"):
            with st.status("Generating Word Document...",expanded=True) as status:
                # Create a Word document from the communication history
                md_text = "\n".join(st.session_state.communication_history) if isinstance(st.session_state.communication_history, list) else str(st.session_state.communication_history)
                doc = convert_markdown_to_word(md_text)

                if doc:
                    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp_docx:
                        doc.save(tmp_docx.name)
                        
                        # Provide the download link
                        st.success("Word Document generated successfully!")
                        st.download_button(label="Download Word Document", data=open(tmp_docx.name, "rb").read(), file_name=f"{st.session_state.first_topic}.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
            status.update(label="Processing complete!", state="complete", expanded=True)  # Finish status
        
        if st.button("Start new chat"):
            del st.session_state["first_topic"]
            del st.session_state["show_input"]
            del st.session_state["communication_history"]
            del st.session_state["generated_details"]
            st.rerun()
